# Copyright (c) 2020 PaddlePaddle Authors. All Rights Reserved.
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

import os
from copy import deepcopy
from math import floor
from shapely.geometry import Polygon

from docx import Document
from docx import shared
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION_START
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.dml.color import ColorFormat, MSO_COLOR_TYPE
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.enum.style import WD_STYLE_TYPE

from ppstructure.recovery.table_process import HtmlToDocx

from ppocr.utils.logging import get_logger


logger = get_logger()


def convert_info_docx(img, res, save_folder, img_name):
    doc = Document()
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    doc.styles['Normal'].font.size = shared.Pt(6.5)
    font_name_title = '宋体'

    flag = 1
    last_bbox_polygon = Polygon([(0, 0), (0, 0), (0, 0), (0, 0)])
    for region in res:
        if len(region['res']) == 0:
            continue
        left, top, right, bottom = region['bbox']
        current_bbox_polygon = Polygon([(left, top), (right, top), (right, bottom), (left, bottom), (left, top)])
        intersection = current_bbox_polygon.intersection(last_bbox_polygon)
        if not intersection.is_empty:
            overlap_area = intersection.area / (current_bbox_polygon.area + last_bbox_polygon.area - intersection.area)
            if overlap_area > 0.8:
                last_bbox_polygon = current_bbox_polygon
                continue
        last_bbox_polygon = current_bbox_polygon

        img_idx = region['img_idx']
        if flag == 2 and region['layout'] == 'single':
            section = doc.add_section(WD_SECTION_START.CONTINUOUS)
            section._sectPr.xpath('./w:cols')[0].set(qn('w:num'), '1')
            flag = 1
        elif flag == 1 and region['layout'] == 'double':
            section = doc.add_section(WD_SECTION_START.CONTINUOUS)
            section._sectPr.xpath('./w:cols')[0].set(qn('w:num'), '2')
            flag = 2

        region_type = region['type'].lower()
        if region_type == 'figure':
            excel_save_folder = os.path.join(save_folder, img_name)
            img_path = os.path.join(excel_save_folder,
                '{}_{}.jpg'.format(region['bbox'], img_idx))
            paragraph_pic = doc.add_paragraph()
            paragraph_pic.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = paragraph_pic.add_run("")
            if flag == 1:
                run.add_picture(img_path, width=shared.Inches(5))
            elif flag == 2:
                run.add_picture(img_path, width=shared.Inches(2))
        elif region_type == 'title':
            paragraph = doc.add_heading('\n'.join([r['text'] for r in region['res']]))
            font = paragraph.style.font
            font.name = font_name_title
            func_set_font = paragraph.style.element.rPr.rFonts.set
            func_set_font(qn('w:eastAsia'), font_name_title)
            func_set_font(qn("w:asciiTheme"), font_name_title)
            func_set_font(qn("w:eastAsiaTheme"), font_name_title)
            font.color.rgb = shared.RGBColor(0x00, 0x00, 0x00)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            font.size = shared.Pt(18)
        elif region_type == 'table':
            parser = HtmlToDocx()
            parser.table_style = 'TableGrid'
            parser.handle_table(region['res']['html'], doc)
        elif region_type == 'header':
            paragraph = doc.add_paragraph(''.join([r['text'] for r in region['res']]))
            paragraph.style.font.size = shared.Pt(14)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif region_type == 'footer':
            section = doc.sections[0]
            footer = section.footer
            footer.add_paragraph(''.join([r['text'] for r in region['res']]))
            # footer.footer_distance = shared.Cm(3)
            # footer.bottom_margin = shared.Cm(5.0)
            footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        else:
            left_top, right_top, right_bottom, left_bottom = region['res'][0]['text_region']
            cn_word_size = (right_bottom[1] + left_bottom[1]) / 2 - (left_top[1] + right_top[1]) / 2
            paragraph = doc.add_paragraph()
            paragraph.style.font.size = shared.Pt(14)

            # Find the very left position x of all lines.
            start_x = 9999
            for line in region['res']:
                x = line['text_region'][0][0]
                if x < start_x:
                    start_x = x

            for line in region['res']:
                x = line['text_region'][0][0]
                indent_word_size = round((x - start_x) / cn_word_size)
                text = line['text']
                if indent_word_size > 0:
                    paragraph.add_run('\n')
                    indent_text = ''.join(['\t' for i in range(floor(indent_word_size / 2))])
                    text = f'{indent_text}{text}'
                paragraph.add_run(text)

    # save to docx
    docx_path = os.path.join(save_folder, '{}_ocr.docx'.format(img_name))
    doc.save(docx_path)
    logger.info('docx save to {}'.format(docx_path))


def sorted_layout_boxes(res, w):
    """
    Sort text boxes in order from top to bottom, left to right
    args:
        res(list):ppstructure results
    return:
        sorted results(list)
    """
    num_boxes = len(res)
    if num_boxes == 1:
        res[0]['layout'] = 'single'
        return res

    sorted_boxes = sorted(res, key=lambda x: (x['bbox'][1], x['bbox'][0]))
    _boxes = list(sorted_boxes)

    new_res = []
    res_left = []
    res_right = []
    i = 0

    while True:
        if i >= num_boxes:
            break
        if i == num_boxes - 1:
            if _boxes[i]['bbox'][1] > _boxes[i - 1]['bbox'][3] and _boxes[i][
                'bbox'][0] < w / 2 and _boxes[i]['bbox'][2] > w / 2:
                new_res += res_left
                new_res += res_right
                _boxes[i]['layout'] = 'single'
                new_res.append(_boxes[i])
            else:
                if _boxes[i]['bbox'][2] > w / 2:
                    _boxes[i]['layout'] = 'double'
                    res_right.append(_boxes[i])
                    new_res += res_left
                    new_res += res_right
                elif _boxes[i]['bbox'][0] < w / 2:
                    _boxes[i]['layout'] = 'double'
                    res_left.append(_boxes[i])
                    new_res += res_left
                    new_res += res_right
            res_left = []
            res_right = []
            break
        elif _boxes[i]['bbox'][0] < w / 4 and _boxes[i]['bbox'][2] < 3 * w / 4:
            _boxes[i]['layout'] = 'double'
            res_left.append(_boxes[i])
            i += 1
        elif _boxes[i]['bbox'][0] > w / 4 and _boxes[i]['bbox'][2] > w / 2:
            _boxes[i]['layout'] = 'double'
            res_right.append(_boxes[i])
            i += 1
        else:
            new_res += res_left
            new_res += res_right
            _boxes[i]['layout'] = 'single'
            new_res.append(_boxes[i])
            res_left = []
            res_right = []
            i += 1
    if res_left:
        new_res += res_left
    if res_right:
        new_res += res_right
    return new_res
